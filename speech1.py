import pyaudio
from subprocess import call
import speech_recognition as sr
import pickle
from docx import Document
from docx.shared import Inches
import os

global start,pic,upload,sub,date,time

def create_file(file):
	doc=Document()
	doc.add_heading(file)
	doc.save(file)
	
def write_to_file(file,msg):
	doc=Document(file)
	doc.add_paragraph(msg)
	doc.save(file)

def take_picture(file):
	read()
	os.system('sudo fswebcam -r 640x420 test.jpg -S 8')
	doc=Document(file)
	doc.add_picture('/home/pi/test.jpg',width=Inches(1.25))
	doc.save(file)

def write(start,pic,upload,sub,date,time):
        v={
                'start':start,
                'pic':pic,
                'upload':upload,
                'sub':sub,
                'date':date,
                'time':time
                }
        with open('var.txt','wb')as handle:
                pickle.dump(v,handle)

def upload_file(filename):
	read()
	os.system('./Dropbox-Uploader/dropbox_uploader.sh upload '+filename+' /'+sub+'/'+filename)
	os.system('sudo rm '+filename)

def read():
	with open('var.txt','rb') as handle:
        	v=pickle.loads(handle.read())
	global start,pic,upload,sub,date,time
	start=v['start']
	pic=v['pic']
	upload=v['upload']
	sub=v['sub']
	date=v['date']
	time=v['time']

def listen(file):
	r = sr.Recognizer()
	r.energy_threshold=4000
	with sr.Microphone(device_index = 2) as source:
		print 'listening..'
		audio = r.listen(source)
		print 'processing'
	try:
    		message = (r.recognize_google(audio, language = 'en-us', show_all=False))
    		write_to_file(file,message)
	except:
    		print('error')


read()
filename=str(sub)+str(date)+str(time)+'.docx'
create_file(filename)

while(upload==False):
        if(pic==True):
		take_picture(filename)
                print('picture taken')
		write(True,False,False,sub,date,time)
		read()
	elif(start==True):
		listen(filename)
		read()
	elif(start==False):
		read()
else:
	upload_file(filename)
	print('file uploaded')
